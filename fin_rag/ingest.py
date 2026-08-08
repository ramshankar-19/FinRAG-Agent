"""
ingest.py — Universal multi-modal ingestion pipeline for Finance RAG.

Supported formats: PDF (text + scanned), DOCX, Images (PNG/JPG/TIFF/BMP)

Pipeline:
    ingest_file(path)
        → detect_file_type()
        → load_*()           [PDF / DOCX / Image]
        → clean_text()
        → chunk
        → embed (all-MiniLM-L6-v2)
        → store in ChromaDB
"""

import os
import re
from pathlib import Path
from typing import List

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings

# ─────────────────────────────────────────────
# Constants
# ─────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {
    ".pdf":  "pdf",
    ".docx": "docx",
    ".doc":  "docx",
    ".png":  "image",
    ".jpg":  "image",
    ".jpeg": "image",
    ".tiff": "image",
    ".tif":  "image",
    ".bmp":  "image",
}

# If avg chars/page falls below this, PDF is treated as scanned → OCR
SCANNED_PDF_THRESHOLD = 100

CHUNK_SIZE    = 1000
CHUNK_OVERLAP = 150
EMBED_MODEL   = "all-MiniLM-L6-v2"

# Standard Tesseract install path on Windows (UB-Mannheim installer)
TESSERACT_WIN_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


# ─────────────────────────────────────────────
# File-type detection
# ─────────────────────────────────────────────

def detect_file_type(file_path: str) -> str:
    """Return 'pdf', 'docx', or 'image' based on file extension."""
    ext = Path(file_path).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Supported extensions: {sorted(SUPPORTED_EXTENSIONS.keys())}"
        )
    return SUPPORTED_EXTENSIONS[ext]


# ─────────────────────────────────────────────
# Text cleaning
# ─────────────────────────────────────────────

def clean_text(text: str) -> str:
    """
    Normalize extracted text:
      - collapse multiple horizontal whitespace into a single space
      - cap consecutive newlines at two
      - strip non-printable control characters (keep \\n and \\t)
    """
    text = re.sub(r"[^\S\n\t]+", " ", text)         # normalise horizontal whitespace
    text = re.sub(r"\n{3,}", "\n\n", text)           # max 2 consecutive newlines
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)  # strip ctrl chars
    return text.strip()


# ─────────────────────────────────────────────
# OCR backends  (modular — swap freely)
# ─────────────────────────────────────────────

def _configure_tesseract() -> bool:
    """
    Point pytesseract at the Tesseract binary if found.
    Returns True if Tesseract is usable, False otherwise.
    """
    try:
        import pytesseract
        import shutil

        binary = shutil.which("tesseract") or (
            TESSERACT_WIN_PATH if os.path.exists(TESSERACT_WIN_PATH) else None
        )
        if binary:
            pytesseract.pytesseract.tesseract_cmd = binary
            return True
        return False
    except ImportError:
        return False


def _ocr_with_pytesseract(image) -> str:
    """
    Primary OCR backend: pytesseract.
    `image` is a PIL.Image instance.
    """
    import pytesseract
    return pytesseract.image_to_string(image)


def _ocr_with_pillow_fallback(image) -> str:
    """
    Fallback OCR: returns a placeholder when no OCR engine is available.
    Preserves the pipeline structure so callers don't need to change.
    """
    return "[OCR unavailable — install Tesseract binary to extract text from this image]"


def run_ocr(image) -> str:
    """
    Dispatch to the best available OCR backend.
    OCR backend is modular: add or replace backends here.
    """
    if _configure_tesseract():
        return _ocr_with_pytesseract(image)
    return _ocr_with_pillow_fallback(image)


# ─────────────────────────────────────────────
# Loaders
# ─────────────────────────────────────────────

def load_pdf(file_path: str) -> List[Document]:
    """
    Load a PDF.
    Text-based PDFs → PyPDFLoader.
    Scanned PDFs (avg chars/page < threshold) → page-by-page OCR.
    """
    from langchain_community.document_loaders import PyPDFLoader

    loader = PyPDFLoader(file_path)
    docs = loader.load()

    if not docs:
        return []

    total_chars = sum(len(d.page_content) for d in docs)
    avg_chars = total_chars / len(docs)

    if avg_chars < SCANNED_PDF_THRESHOLD:
        print(f"  Scanned PDF detected (avg {avg_chars:.0f} chars/page). Applying OCR …")
        return _load_scanned_pdf(file_path, n_pages=len(docs))

    return docs


def _load_scanned_pdf(file_path: str, n_pages: int) -> List[Document]:
    """OCR each page of a scanned PDF using pdf2image + Tesseract."""
    try:
        from pdf2image import convert_from_path
        from PIL import Image

        images = convert_from_path(file_path, dpi=200)
        docs = []
        for i, img in enumerate(images):
            text = run_ocr(img)
            docs.append(Document(
                page_content=text,
                metadata={"source": file_path, "page": i, "ocr": True}
            ))
        return docs
    except ImportError as exc:
        raise RuntimeError(
            "pdf2image is required for scanned PDF OCR. "
            "Install it with: pip install pdf2image"
        ) from exc


def load_docx(file_path: str) -> List[Document]:
    """
    Load a DOCX file using python-docx.
    Extracts paragraphs and table rows as a single Document.
    """
    import docx

    doc = docx.Document(file_path)
    parts: List[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            parts.append(stripped)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    full_text = "\n".join(parts)
    return [Document(
        page_content=full_text,
        metadata={"source": file_path}
    )]


def load_image(file_path: str) -> List[Document]:
    """
    Load an image file and extract text via OCR.
    OCR backend is modular (see run_ocr).
    """
    from PIL import Image

    img = Image.open(file_path)
    text = run_ocr(img)

    if not text.strip():
        text = "[No text could be extracted from this image]"

    return [Document(
        page_content=text,
        metadata={"source": file_path, "ocr": True}
    )]


# ─────────────────────────────────────────────
# Main public function
# ─────────────────────────────────────────────

def ingest_file(
    file_path: str,
    persist_directory: str = "data/chroma_db",
) -> int:
    """
    Universal ingestion pipeline.

    Steps:
        1. Detect file type
        2. Load / OCR
        3. Clean text
        4. Chunk
        5. Embed (all-MiniLM-L6-v2)
        6. Store in ChromaDB

    Returns:
        Number of chunks stored.

    Raises:
        FileNotFoundError  — file_path does not exist
        ValueError         — unsupported file extension
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    file_name = Path(file_path).name
    doc_type  = detect_file_type(file_path)

    print(f"[Ingest] {file_name}  (type={doc_type})")

    # ── Step 1: Load ──────────────────────────────────
    print("  → Loading …")
    if doc_type == "pdf":
        raw_docs = load_pdf(file_path)
    elif doc_type == "docx":
        raw_docs = load_docx(file_path)
    elif doc_type == "image":
        raw_docs = load_image(file_path)
    else:
        raise ValueError(f"Unhandled doc_type: {doc_type}")

    print(f"  → Loaded {len(raw_docs)} document section(s)")

    # ── Step 2: Clean + tag metadata ─────────────────
    processed: List[Document] = []
    for idx, doc in enumerate(raw_docs):
        cleaned = clean_text(doc.page_content)
        if len(cleaned) < 20:       # skip near-empty pages
            continue
        doc.page_content = cleaned
        doc.metadata.update({
            "file_name": file_name,
            "doc_type":  doc_type,
        })
        processed.append(doc)

    if not processed:
        print("  ⚠ No usable text extracted — nothing stored.")
        return 0

    # ── Step 3: Chunk ─────────────────────────────────
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )
    chunks = splitter.split_documents(processed)

    for i, chunk in enumerate(chunks):
        chunk.metadata["chunk_index"] = i

    print(f"  → Split into {len(chunks)} chunks")

    # ── Step 4: Embed + Store ─────────────────────────
    print("  → Embedding and storing …")
    embeddings = HuggingFaceEmbeddings(model_name=EMBED_MODEL)
    os.makedirs(persist_directory, exist_ok=True)

    vectorstore = Chroma(
        persist_directory=persist_directory,
        embedding_function=embeddings,
    )
    vectorstore.add_documents(chunks)

    print(f"  ✓ Stored {len(chunks)} chunks in ChromaDB → '{persist_directory}'")
    return len(chunks)


# ─────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python ingest.py <file_path> [chroma_persist_dir]")
        print("Supported: .pdf  .docx  .doc  .png  .jpg  .jpeg  .tiff  .bmp")
        sys.exit(1)

    path = sys.argv[1]
    db   = sys.argv[2] if len(sys.argv) > 2 else "data/chroma_db"

    os.makedirs("data", exist_ok=True)
    n = ingest_file(path, persist_directory=db)
    print(f"\nDone. {n} chunks indexed into '{db}'.")